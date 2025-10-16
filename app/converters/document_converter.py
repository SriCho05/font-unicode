"""
Document converter module for processing various file formats
"""
import os
import tempfile
import shutil
from pathlib import Path
import chardet

# Optional imports with fallbacks
try:
    import magic
except ImportError:
    magic = None

# Document processing libraries
try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    Document = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import pandas as pd
except ImportError:
    pd = None

class DocumentConverter:
    def __init__(self, font_detector, font_mapper):
        self.font_detector = font_detector
        self.font_mapper = font_mapper
        
        # Supported file formats
        self.supported_formats = {
            'txt': self._convert_txt,
            'docx': self._convert_docx,
            'doc': self._convert_doc,
            'pdf': self._convert_pdf
        }
    
    def convert_document(self, file_path):
        """
        Convert a document from non-Unicode to Unicode fonts
        
        Args:
            file_path (str): Path to the input document
            
        Returns:
            dict: Conversion result with success status, output file, and statistics
        """
        try:
            # Detect file format
            file_extension = Path(file_path).suffix.lower().lstrip('.')
            
            if file_extension not in self.supported_formats:
                return {
                    'success': False,
                    'error': f'Unsupported file format: {file_extension}'
                }
            
            # Get file info
            file_info = self._get_file_info(file_path)
            
            # Convert based on file type
            converter_func = self.supported_formats[file_extension]
            result = converter_func(file_path)
            
            # Add file info to result
            result['file_info'] = file_info
            
            return result
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error converting document: {str(e)}'
            }
    
    def _get_file_info(self, file_path):
        """Get basic file information"""
        stat = os.stat(file_path)
        return {
            'name': os.path.basename(file_path),
            'size': stat.st_size,
            'extension': Path(file_path).suffix.lower()
        }
    
    def _convert_txt(self, file_path):
        """Convert plain text file"""
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                encoding_result = chardet.detect(raw_data)
                encoding = encoding_result.get('encoding', 'utf-8')
            
            # Read text content
            with open(file_path, 'r', encoding=encoding) as f:
                original_text = f.read()
            
            # Detect fonts
            detection_result = self.font_detector.detect_fonts(original_text)
            
            # Convert text
            converted_text = self.font_mapper.convert_with_preservation(original_text)
            
            # Generate output filename
            output_filename = f"converted_{os.path.basename(file_path)}"
            output_path = os.path.join('app/downloads', output_filename)
            
            # Save converted text
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(converted_text)
            
            # Generate statistics
            stats = self.font_mapper.get_conversion_stats(original_text, converted_text)
            stats['detected_fonts'] = detection_result
            
            return {
                'success': True,
                'output_filename': output_filename,
                'output_path': output_path,
                'preview': {
                    'original': original_text[:500] + '...' if len(original_text) > 500 else original_text,
                    'converted': converted_text[:500] + '...' if len(converted_text) > 500 else converted_text
                },
                'stats': stats
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error converting TXT file: {str(e)}'
            }
    
    def _convert_docx(self, file_path):
        """Convert DOCX file while preserving formatting"""
        if not Document:
            return {
                'success': False,
                'error': 'python-docx library not available for DOCX processing'
            }
        
        try:
            # Load document
            doc = Document(file_path)
            
            original_text = ""
            converted_text = ""
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                original_para_text = paragraph.text
                original_text += original_para_text + "\n"
                
                if original_para_text.strip():
                    # Convert text while preserving formatting
                    converted_para_text = self.font_mapper.convert_with_preservation(original_para_text)
                    converted_text += converted_para_text + "\n"
                    
                    # Update paragraph text
                    paragraph.text = converted_para_text
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        original_cell_text = cell.text
                        if original_cell_text.strip():
                            converted_cell_text = self.font_mapper.convert_with_preservation(original_cell_text)
                            cell.text = converted_cell_text
            
            # Generate output filename
            output_filename = f"converted_{os.path.basename(file_path)}"
            output_path = os.path.join('app/downloads', output_filename)
            
            # Save converted document
            doc.save(output_path)
            
            # Detect fonts and generate statistics
            detection_result = self.font_detector.detect_fonts(original_text)
            stats = self.font_mapper.get_conversion_stats(original_text, converted_text)
            stats['detected_fonts'] = detection_result
            
            return {
                'success': True,
                'output_filename': output_filename,
                'output_path': output_path,
                'preview': {
                    'original': original_text[:500] + '...' if len(original_text) > 500 else original_text,
                    'converted': converted_text[:500] + '...' if len(converted_text) > 500 else converted_text
                },
                'stats': stats
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error converting DOCX file: {str(e)}'
            }
    
    def _convert_doc(self, file_path):
        """Convert DOC file (legacy Word format)"""
        # For DOC files, we'll need to use a different approach
        # This is a simplified implementation
        return {
            'success': False,
            'error': 'DOC file conversion not fully implemented. Please convert to DOCX format first.'
        }
    
    def _convert_pdf(self, file_path):
        """Convert PDF file"""
        if not PyPDF2:
            return {
                'success': False,
                'error': 'PyPDF2 library not available for PDF processing'
            }
        
        try:
            # Read PDF content
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                original_text = ""
                for page in pdf_reader.pages:
                    original_text += page.extract_text() + "\n"
            
            # Convert text
            converted_text = self.font_mapper.convert_with_preservation(original_text)
            
            # Generate output filename (as text file since PDF editing is complex)
            base_name = Path(file_path).stem
            output_filename = f"converted_{base_name}.txt"
            output_path = os.path.join('app/downloads', output_filename)
            
            # Save converted text
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(converted_text)
            
            # Generate statistics
            detection_result = self.font_detector.detect_fonts(original_text)
            stats = self.font_mapper.get_conversion_stats(original_text, converted_text)
            stats['detected_fonts'] = detection_result
            stats['note'] = 'PDF converted to text format due to formatting complexity'
            
            return {
                'success': True,
                'output_filename': output_filename,
                'output_path': output_path,
                'preview': {
                    'original': original_text[:500] + '...' if len(original_text) > 500 else original_text,
                    'converted': converted_text[:500] + '...' if len(converted_text) > 500 else converted_text
                },
                'stats': stats
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error converting PDF file: {str(e)}'
            }
    
    def get_supported_formats(self):
        """Get list of supported file formats"""
        return list(self.supported_formats.keys())
    
    def validate_file(self, file_path):
        """Validate if file can be processed"""
        if not os.path.exists(file_path):
            return False, "File does not exist"
        
        file_extension = Path(file_path).suffix.lower().lstrip('.')
        if file_extension not in self.supported_formats:
            return False, f"Unsupported file format: {file_extension}"
        
        # Check file size (max 16MB)
        file_size = os.path.getsize(file_path)
        if file_size > 16 * 1024 * 1024:
            return False, "File too large (max 16MB)"
        
        return True, "File is valid"